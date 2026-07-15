#!/usr/bin/env python3
"""Generate weekly progress report Word document (PR4) aligned with course rubric."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).resolve().parent
OUT_PATH = OUT_DIR / "PR4_Weekly_Progress_Report.docx"

TEAM = [
    ("Afolabi Adesina", "991808102"),
    ("Evans Frimpong", "991771602"),
    ("Xian Qin", "991578381"),
    ("Yanan Yang", "991841037"),
    ("", ""),
]

GITHUB = "https://github.com/afolabi-adesina-m/Smart-Shield-AI"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _rubric_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI/ML Capstone Project\nPR4 – Weekly Progress Report (5%)")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Instructor: Prof. Mouhamed Abdulla  |  Course: INFO53883  |  Team 2B\n"
        f"Sprint: RW/W8 (Jun 23 – Jun 30, 2026)  |  Report date: {date.today().strftime('%B %d, %Y')}"
    )

    doc.add_paragraph()
    for i, (name, sid) in enumerate(TEAM, 1):
        if name:
            doc.add_paragraph(f"Name {i}: {name}    Student Number {i}: {sid}")

    doc.add_page_break()

    # Q1
    _heading(doc, "Q1: Summary of Work", 2)
    _rubric_note(doc, "Rubric: Provide an overview of completed work during this sprint (~⅓ page, single-space). [1 mark]")
    doc.add_paragraph(
        "This sprint focused on completing and documenting exploratory data analysis (EDA), "
        "aligning findings with our two core research papers, and hardening the project for team collaboration."
    )

    _heading(doc, "A. Exploratory Data Analysis — Toronto Police Service (TPS)", 3)
    _bullets(
        doc,
        [
            "Loaded df_toronto: 809,034 rows × 21 columns from traffic collision data.csv (Toronto Open Data).",
            "Engineered 3-class SEVERITY target: 0 = property damage only (86.36%), 1 = injury (13.56%), 2 = fatal (0.08%).",
            "Temporal EDA (6-panel plots): rush-hour peaks (morning/evening), seasonal variation, and division/neighbourhood concentration — supports OCC_HOUR, IS_NIGHT, IS_RUSHHOUR, and SEASON_NUM as model inputs.",
            "After preprocessing: 809,030 clean rows; 677,000+ records with valid GPS for map-based demo routing.",
            "Correlation analysis: pedestrian involvement (r ≈ 0.32) and bicycle involvement (r ≈ 0.22) show the strongest linear association with severity; all eight candidate features significant at p < 0.05 (chi-square).",
            "Three-method feature voting (chi², mutual information, Random Forest Gini): all 8 features retained — OCC_HOUR, MONTH_NUM, SEASON_NUM, IS_NIGHT, IS_RUSHHOUR, PEDESTRIAN_BIN, BICYCLE_BIN, AUTOMOBILE_BIN.",
        ],
    )

    _heading(doc, "B. Exploratory Data Analysis — UK DfT 2024", 3)
    _bullets(
        doc,
        [
            "Loaded dft: 100,927 rows × 44 columns (dft-road-casualty-statistics-collision-2024.csv).",
            "Decoded road-surface codes: wet/damp, frost/ice, snow, and flood conditions show higher serious/fatal proportions than dry surfaces (e.g., wet/damp ~24.1% serious vs dry ~23.6%).",
            "Replicated Paper 2 summary tables (weather, road surface, lighting): fog/mist and rain+wind combinations yield highest mean casualties per collision on our DfT subset.",
            "Excess-casualty analysis: fog/mist +15.8% above baseline; flood and snow/slush also exceed mean — used to calibrate E_index weights for the Safety Score.",
            "Road surface × weather heatmap: snow/slush combined with snowing+wind identified as the highest-risk pairing — informs Ontario 400-series winter demo presets.",
            "Point-biserial / Cramér's V: lighting and speed limit correlate with severe outcomes; weather/surface codes work best as composite environmental indices rather than single predictors.",
        ],
    )

    _heading(doc, "C. Literature Papers Used & Key Findings Applied", 3)
    _bullets(
        doc,
        [
            "Paper 1 — Pennino & D'Amato (2024), Enhancing Safety in Autonomous Navigation: Weather-Aware Trajectory Planning (IEEE MetroSea). "
            "Composite Seakeeping Performance Index (SPI) mapped to our Safety Score S (0–100) with LOW / MEDIUM / HIGH risk tiers and speed-reduction advisories.",
            "Paper 2 — Jiang, Miao & Wu (2024), Machine Learning based Prediction Analysis of Traffic Accidents. "
            "Benchmark weather/surface/lighting severity tables replicated on UK DfT data; chi-square confirms weather, road condition, lighting, and junction type as significant predictors (all p < 0.01 in published results).",
            "E_index weight calibration from Paper 2 excess casualties: surface hazard δ=0.35, visibility γ=0.30, wind/precipitation β=0.20, temperature α=0.15.",
            "Integrated both papers into Section 0b of capstone_with_results.ipynb with annotated markdown (What this cell does / Findings & importance) for presentation-ready documentation.",
        ],
    )

    _heading(doc, "D. Additional Sprint Deliverables", 3)
    _bullets(
        doc,
        [
            "Repository reorganization (Data/, notebooks/, src/, models/, demo/) with auto-detecting DATA path and SMART_SHIELD_DATA environment override.",
            "Annotated notebook (capstone_with_results.ipynb) with saved EDA outputs, collapsible sections, and post-run documentation pipeline.",
            "explanations/ auto-generated glossary, sprint swimlane, and train/test split reference docs.",
            "Continued Vision Brain cache setup and Flask map demo (api_server.py) loading trained tabular artifacts.",
        ],
    )

    # Q2
    doc.add_paragraph()
    _heading(doc, "Q2: List of Challenges", 2)
    _rubric_note(doc, "Rubric: List challenges and solutions encountered (~⅓ page, single-space). [1 mark]")
    _bullets(
        doc,
        [
            "Challenge: Team members could not locate CSV files after June 2026 folder reorganization (data appeared missing). "
            "Solution: Standardized Data/ folder, updated notebook _resolve_data_dir(), documented filenames in Data/README.md, and recommended shared OneDrive/Teams folder plus SMART_SHIELD_DATA env var.",
            "Challenge: Extreme class imbalance on Toronto SEVERITY (86% majority class). "
            "Solution: SMOTE on training split only, Matthews Correlation Coefficient (MCC) as primary metric, stratified 80/20 split and 3-fold CV in GridSearchCV.",
            "Challenge: Paper 2 replication required harmonized DfT code labels across weather, surface, and lighting fields. "
            "Solution: Built decode maps and filtered subsets matching published table definitions before plotting Tables A/B/C.",
            "Challenge: Collaborators editing different notebook copies caused confusion. "
            "Solution: Designated capstone.ipynb as working notebook, capstone_with_results.ipynb as annotated output; branch protection and PR workflow recommended on GitHub.",
            "Challenge: Vision training limited by small labelled Ontario road-image cache. "
            "Solution: ResNet18 transfer learning from ImageNet, synthetic/demo images in vision_cache/, and documented HuggingFace RSCD as future scale-up source.",
        ],
    )

    # Q3
    doc.add_paragraph()
    _heading(doc, "Q3: Technology Stack", 2)
    _rubric_note(doc, "Rubric: Tools, libraries, DBs, APIs with version numbers (~⅓ page, single-space). [1 mark]")
    _bullets(
        doc,
        [
            "IDE / notebooks: Cursor, Visual Studio Code, Jupyter Notebook; Python 3.10 (conda env: ai_work_final).",
            "Core Python: pandas 2.3.3, NumPy, SciPy, matplotlib, seaborn.",
            "Machine learning: scikit-learn 1.7.2, LightGBM 4.6.0, imbalanced-learn (SMOTE), SHAP.",
            "Deep learning: PyTorch 2.12.1+cpu, torchvision, ResNet18 fine-tuning for Vision Brain.",
            "Model persistence: joblib (.joblib), PyTorch state dict (.pt).",
            "Web demo: Flask, flask-cors, python-dotenv; OpenStreetMap + OSRM routing API (no API key required).",
            "NLP pillar: TF-IDF vectorizer (scikit-learn) on Ontario 511-style alert text.",
            "Data sources: Toronto Police Service Open Data (traffic collision data.csv); UK DfT Road Safety Data (collision-2024.csv); local vision_cache/ JPEG images.",
            "Version control: Git, GitHub (https://github.com/afolabi-adesina-m/Smart-Shield-AI).",
            "Document generation: python-docx, nbformat, nbclient.",
        ],
    )

    # Q4
    doc.add_paragraph()
    _heading(doc, "Q4: Meeting Minutes & GitHub Activity", 2)
    _rubric_note(doc, "Rubric: Key highlights from in-class and outside meetings; GitHub URLs (~⅔ page, single-space). [1 mark]")
    _bullets(
        doc,
        [
            "In-class (W8): Reviewed EDA plots and Paper 2 replication tables; professor feedback to emphasize MCC over accuracy for imbalanced severity classes.",
            "Team sync (Jun 26): Agreed canonical data location (Data/ at repo root); Afolabi demonstrated notebook auto-path resolution after reorg.",
            "Team sync (Jun 28): Reviewed annotated capstone_with_results.ipynb for PR submission; assigned Evans to demo API integration, Yanan to imbalance experiments, Xian to weather API feeds.",
            "Decision: Publish code/notebooks on GitHub; keep large CSVs and model weights in shared Drive (gitignored) to prevent teammate data-loss issues.",
            "Decision: Use 3-Brain architecture (NLP + Vision + Tabular) with fused Safety Score S for final presentation.",
            f"GitHub — repository: {GITHUB}",
            f"GitHub — latest commit (reorg + mobile demo): {GITHUB}/commits/main",
            f"GitHub — notebooks: {GITHUB}/tree/main/notebooks",
            f"GitHub — demo app: {GITHUB}/tree/main/demo",
            f"GitHub — explanations docs: {GITHUB}/tree/main/explanations",
        ],
    )

    # Q5
    doc.add_paragraph()
    _heading(doc, "Q5: Future Work (Upcoming Week — W9)", 2)
    _rubric_note(doc, "Rubric: Work plan for the upcoming week (~⅓ page, single-space). [1 mark]")
    _bullets(
        doc,
        [
            "Complete multimodal fusion (Section 10): combine TF-IDF text score (T), Vision V-score, and tabular/environment E_index into final Safety Score S.",
            "Finalize tuned Random Forest / LightGBM deployment artifacts and run held-out test evaluation with confusion matrices for PR5.",
            "Evans: Integrate demo map with live scoring API and weather presets (clear, wet, blizzard, ice storm).",
            "Yanan: Run SMOTE + class-weight experiments; report macro-F1 and recall for fatal/injury classes.",
            "Xian: Validate Environment Canada / Ontario 511 API ingestion for real-time E_index inputs.",
            "Afolabi: SHAP explainability cells, speed-advisory audit fixes (SS-AUDIT-2026-001), and PR5 draft.",
            "Team: Upload PR4 PDF to SLATE; staple hardcopy for in-class submission (deadline Wed Jul 1, 2026 per course schedule).",
        ],
    )

    # References
    doc.add_page_break()
    _heading(doc, "References", 2)
    refs = [
        '[1] F. Pennino and A. D\'Amato, "Enhancing Safety in Autonomous Navigation: Weather-Aware Trajectory Planning," in Proc. IEEE MetroSea, 2024. (Sheridan Library – IEEE Xplore).',
        '[2] J. Jiang, Y. Miao, and D. Wu, "Machine Learning based Prediction Analysis of Potential Factors in Traffic Accidents," Applied and Computational Engineering, vol. 99, pp. 112–120, Nov. 2024. https://doi.org/10.54254/2755-2721/99/20251788',
        '[3] Toronto Police Service, "Traffic collision data," City of Toronto Open Data Portal. https://open.toronto.ca/',
        '[4] Department for Transport (UK), "Road casualty statistics, collisions 2024," 2025. https://data.dft.gov.uk/road-accidents-safety-data/',
        '[5] scikit-learn developers, "scikit-learn 1.7.2 documentation," 2025. https://scikit-learn.org/stable/',
        '[6] PyTorch Team, "PyTorch 2.12 documentation," 2025. https://pytorch.org/docs/',
        '[7] Microsoft, "Visual Studio Code documentation," 2025. https://code.visualstudio.com/docs',
        '[8] Cursor, "Cursor IDE documentation," 2025. https://cursor.com/docs',
        '[9] Smart-Shield AI Capstone Repository (Team 2B), 2026. ' + GITHUB,
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    doc.add_paragraph()
    footer = doc.add_paragraph("Sheridan College — School of Applied Computing — INFO53883 AI/ML Capstone — Spring/Summer 2026")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Created: {path}")
