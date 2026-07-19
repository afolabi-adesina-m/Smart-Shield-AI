#!/usr/bin/env python3
"""Generate This Week's Project Report (PR5 / W9) — Text, Vision, Environment pillars."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).resolve().parent
OUT_PATH = OUT_DIR / "This_Weeks_Project_Report.docx"

TEAM = [
    ("Afolabi Adesina", "991808102"),
    ("Evans Frimpong", "991771602"),
    ("Xian Qin", "991578381"),
    ("Yanan Yang", "991841037"),
]

GITHUB = "https://github.com/afolabi-adesina-m/Smart-Shield-AI"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _rubric_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(10)


def _para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "AI/ML Capstone Project\nThis Week's Project Report (PR5 – Weekly Progress Report 5%)"
    )
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Instructor: Prof. Mouhamed Abdulla  |  Course: INFO53883  |  Team 2B\n"
        f"Sprint: W9 (Jul 1 – Jul 7, 2026)  |  Report date: {date.today().strftime('%B %d, %Y')}"
    )

    doc.add_paragraph()
    for i, (name, sid) in enumerate(TEAM, 1):
        doc.add_paragraph(f"Name {i}: {name}    Student Number {i}: {sid}")

    doc.add_page_break()

    # ── Q1 ─────────────────────────────────────────────────────────────────────
    _heading(doc, "Q1: Summary of Work", 2)
    _rubric_note(
        doc,
        "Rubric: Provide an overview of completed work during this sprint (~⅓ page, single-space). [1 mark]",
    )

    _heading(doc, "Project Objective (Clarified Scope)", 3)
    _para(
        doc,
        "Smart-Shield AI scores one Ontario highway route per search (not all possible routes). "
        "For each trip, the system fuses three real-time risk signals into a Safety Score S (0–100) "
        "and returns a speed advisory (e.g., 80 km/h on Hwy 400 in a blizzard vs 100 km/h on a clear day). "
        "Training data for the tabular collision model is Toronto TPS (~809,034 records; ~768,000+ after cleaning). "
        "UK DfT and Seattle SDOT are used to replicate Paper 2 environmental statistics — not to train the Ontario severity classifier.",
    )
    _para(
        doc,
        "Important distinction: Text + Vision produce real-time hazard signals (T and V). "
        "Environment produces the weather/time/surface risk index (E). "
        "Tabular Random Forest models historical collision severity on Toronto separately from T/V/E fusion.",
    )

    _heading(doc, "Pillar 1 — Text Brain (NLP) · 25% of Safety Score S", 3)
    _bullets(
        doc,
        [
            "Module: src/nlp_brain.py  |  Notebook: Section 10.1 (capstone_with_results.ipynb)",
            "Method: TF-IDF (ngram_range 1–2) + 16-term Ontario 511 hazard lexicon (ice, blizzard, black ice, multi-vehicle, etc.)",
            "Output: T ∈ [0, 1] — normalized hazard score from alert text",
            "Saved artifact: models/tfidf_vectorizer.joblib",
            "Validated on 5 scenarios: TC-1 T=0.000 | TC-2 T=0.688 | TC-3 T=0.000 | TC-4 T=0.000 | TC-5 T=1.000",
            "Demo: demo/inference.py maps weather presets to Ontario 511-style alert strings",
            "Role: Scores live alert language for fusion — does NOT classify tabular collision severity.",
        ],
    )

    _heading(doc, "Pillar 2 — Vision Brain (Computer Vision) · 35% of Safety Score S", 3)
    _bullets(
        doc,
        [
            "Module: src/vision_brain.py  |  Notebook: Sections 6.2, 6.2b, 6.2c",
            "Architecture: ResNet18 transfer learning → 4 road-surface classes (Clear Asphalt, Wet/Slush, Snow/Ice)",
            "Hybrid (6.2b): Conv autoencoder on clear roads → V_anomaly; fused: V_vision = 0.70·V_class + 0.30·V_anomaly",
            "Analysis (6.2c): Softmax probabilities, latent t-SNE, per-class hazard summaries",
            "Output: V ∈ [0, 1] — surface/anomaly risk (not collision severity from tabular data)",
            "Training: Val accuracy 100% on 72 hold-out images (288 train; small cache — documented limitation)",
            "Saved artifact: models/vision_resnet18.pt (deployed Jul 6, 2026)",
            "Demo presets when hybrid not trained: clear=0.08, wet=0.42, blizzard=0.88, ice_storm=0.82",
            "Status: ResNet18 already implemented; YOLO not required for surface classification.",
        ],
    )

    _heading(doc, "Pillar 3 — Environment Brain (E_index) · 40% of Safety Score S", 3)
    _bullets(
        doc,
        [
            "Module: src/safety_score.py → e_index_from_features(), compute_e_index()",
            "Notebook: Sections 2.4, 2.5, 9.5, 10.2",
            "Formula: E = 0.35·surface + 0.30·visibility + 0.20·wind + 0.15·temp (Paper 2 calibrated weights)",
            "Inputs: month_num, season_num, is_night, is_winter_storm from TC scenario features",
            "Paper 2 replication: df_paper2_merged = 296,250 rows (SDOT 203,367 + UK DfT 92,883)",
            "EDA alignment: Dusk +14.77% excess casualties (paper: +13.01%); Snow/Slush road +7.89% (paper: +10.75%)",
            "Live outputs: TC-2 blizzard night E=0.940 | TC-1 clear summer E=0.165",
        ],
    )

    _heading(doc, "Tabular Model (Toronto TPS — Supports Historical Risk Context)", 3)
    _bullets(
        doc,
        [
            "Dataset: 3-class SEVERITY — PD-only 86.36%, Injury 13.56%, Fatal 0.08%",
            "SMOTE on training set only (~1.67M after resample, 33% per class); test set keeps real-world imbalance",
            "Five baseline models trained: best accuracy kNN 88.22% (macro-F1 42.78%, MCC 0.3605)",
            "Deployed model: Random Forest (tuned) — best MCC + interpretability balance",
            "Honest limitation: Fatal class F1 ≈ 0 across all models — documented in Section 7.8",
            "Per-class classification report and inline KPI gate created (Yanan)",
        ],
    )

    _heading(doc, "Fusion & Deployment (All Three Pillars Combined)", 3)
    _bullets(
        doc,
        [
            "Formula: S = (0.25·T + 0.35·V + 0.40·E) × 100",
            "Notebook: Section 9.5 (unseen-data trace per pillar), Section 10.2 (fusion dashboard)",
            "Example: TC-2 Blizzard S=87.0 HIGH 80 km/h | TC-1 Clear S=11.9 LOW 100 km/h",
            "API/Demo: demo/api_server.py (desktop :5050), demo/mobile_server.py (:5051) — OSRM + fused scoring",
            "Models saved: rf_tuned.joblib, lr_tuned.joblib, dnn_smart_shield.pt, vision_resnet18.pt, tfidf_vectorizer.joblib, scaler.joblib",
        ],
    )

    _heading(doc, "Team Contributions (This Sprint)", 3)
    _bullets(
        doc,
        [
            "Afolabi Adesina: Hybrid Vision Brain (ResNet18 + autoencoder), Sections 6.2b/6.2c/7.8/9.5/10.2, demo/inference.py fusion wiring, primary notebook capstone_with_results.ipynb",
            "Yanan Yang: Reorganized data inspection, EDA, preprocessing, visualization on GitHub; SDOT severity engineering in fixed_Y_V1; five baseline models, SMOTE, per-class KPI gate; three-dataset strategy (Toronto train, SDOT+DfT Paper 2 replication)",
            "Evans Frimpong: Repository restructure (Data/, notebooks/, src/, models/, demo/, explanations/); Flask map demo + mobile view; GitHub-ready layout (commit 614ea7e, Jul 6)",
            "Xian Qin: Environment Canada / Ontario 511 data source research for live E_index feeds (in progress)",
        ],
    )

    # ── Q2 ─────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "Q2: List of Challenges", 2)
    _rubric_note(
        doc,
        "Rubric: List challenges and solutions encountered (~⅓ page, single-space). [1 mark]",
    )
    _bullets(
        doc,
        [
            "Fatal class F1 ≈ 0 on all tabular models (Tabular). "
            "Solution: Documented in Section 7.8; primary metric shifted to MCC; real-time risk uses T+V+E fusion, not RF fatal prediction alone.",
            "SMOTE cannot fix fatal separation (Tabular). "
            "Solution: Training balanced to 33% per class but test stays imbalanced; proves need for multimodal fusion beyond tabular features.",
            "Small vision cache — 17 real images (Vision). "
            "Solution: ResNet18 transfer learning + synthetic augmentation; 100% val accuracy flagged as overfit risk in notebook.",
            "Two SDOT severity paths in fixed_Y_V1 (Environment). "
            "Solution: Flag-based assign_sdot_severity vs SEVERITYCODE remap conflict identified; main notebook uses harmonised Paper 2 merge.",
            "CPU training timeouts during vision fine-tuning (Vision). "
            "Solution: Documented skip path — run 6.2 only, skip 6.2b, use V presets in fusion.",
            "Paper 2 ML accuracy gap — ours 0.75 vs paper 0.878 (Tabular). "
            "Solution: Section 7.8 explains task mismatch — Toronto 3-class vs paper binary on SDOT+UK.",
            "Dataset limitations — hour, season, pedestrian flags cannot separate fatal crashes (Tabular). "
            "Solution: Advanced GridSearchCV tuning planned (Yanan); fusion architecture compensates for tabular ceiling.",
        ],
    )

    # ── Q3 ─────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "Q3: Technology Stack", 2)
    _rubric_note(
        doc,
        "Rubric: Tools, libraries, DBs, APIs with version numbers (~⅓ page, single-space). [1 mark]",
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Tool / Library / Dataset"
    hdr[2].text = "Version"
    rows = [
        ("Programming Language", "Python (conda env: ai_work_final)", "3.10"),
        ("Data Manipulation", "pandas, NumPy, SciPy", "2.3.3 / latest"),
        ("Machine Learning & Tuning", "scikit-learn", "1.7.2"),
        ("Gradient Boosting", "LightGBM", "4.6.0"),
        ("Deep Learning", "PyTorch, torchvision", "2.12.1+cpu"),
        ("Data Resampling", "imbalanced-learn (SMOTE)", "0.11+"),
        ("Visualization", "matplotlib, seaborn", "3.7+"),
        ("Model Explainability", "SHAP", "0.42+"),
        ("NLP — Text Pillar", "TF-IDF (scikit-learn)", "in sklearn 1.7.2"),
        ("Vision — Vision Pillar", "ResNet18 + custom autoencoder", "torchvision"),
        ("Web Demo", "Flask, flask-cors, OSRM, Leaflet", "latest"),
        ("Model Persistence", "joblib, PyTorch .pt", "1.3+"),
        ("Toronto Training Data", "Toronto TPS collision data", "~809,034 rows"),
        ("Paper 2 Reference Data", "UK DfT 2024 + Seattle SDOT", "100,927 + 260,312 rows"),
        ("Version Control", "Git, GitHub", GITHUB),
        ("Document Generation", "python-docx", "latest"),
    ]
    for comp, tool, ver in rows:
        row = table.add_row().cells
        row[0].text = comp
        row[1].text = tool
        row[2].text = ver

    # ── Q4 ─────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "Q4: Meeting Minutes & GitHub Activity", 2)
    _rubric_note(
        doc,
        "Rubric: Key highlights from in-class and outside meetings; GitHub URLs (~⅔ page, single-space). [1 mark]",
    )
    _bullets(
        doc,
        [
            "In-class (W9): Every team member walked through last week's progress; joint code review on notebook Sections 6–10.",
            "Team sync (Jul 6): Committed SDOT integration, data quality checks, annotated notebook, explanations/ docs.",
            "Decision: capstone_with_results.ipynb is the canonical submission notebook.",
            "Decision: Score one route per search with real-time T + V + E; tabular RF supports historical context, not sole brain.",
            "Decision: Publish code/notebooks on GitHub; large CSVs in shared Drive (gitignored).",
            f"GitHub — repository: {GITHUB}",
            f"GitHub — latest commit (Jul 6): {GITHUB}/commits/master",
            f"GitHub — notebooks: {GITHUB}/tree/main/notebooks",
            f"GitHub — demo app: {GITHUB}/tree/main/demo",
            f"GitHub — models: {GITHUB}/tree/main/models",
            f"GitHub — explanations docs: {GITHUB}/tree/main/explanations",
        ],
    )

    # ── Q5 ─────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "Q5: Future Work (Upcoming Week — W10)", 2)
    _rubric_note(
        doc,
        "Rubric: Work plan for the upcoming week (~⅓ page, single-space). [1 mark]",
    )
    _bullets(
        doc,
        [
            "Yanan: GridSearchCV advanced tuning on all five tabular models; report MCC + per-class fatal recall; ensure I/O compatible with backend.",
            "Yanan: Fix SDOT severity merge in fixed_Y_V1 (use flag-based SEVERITY consistently in merge).",
            "Xian: Live Environment Canada XML feed → E_index (replace static month/season flags).",
            "Afolabi: Run Section 6.2b on GPU/cloud; deploy vision_autoencoder.pt + vision_meta.json; final PR PDF.",
            "Evans: Final presentation demo — Toronto → Barrie, clear vs ice storm side-by-side on map.",
            "Team: Upload This Week's Project Report to SLATE; prep final presentation.",
        ],
    )

    _para(
        doc,
        "Already completed (do not list as future work): ResNet18 vision pipeline, Flask API endpoints, "
        "repository restructure, multimodal fusion Section 10.2, SHAP explainability, model deployment artifacts.",
    )

    # ── References ───────────────────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "References", 2)
    refs = [
        '[1] F. Pennino and A. D\'Amato, "Enhancing Safety in Autonomous Navigation: Weather-Aware Trajectory Planning," in Proc. IEEE MetroSea, 2024.',
        '[2] J. Jiang, Y. Miao, and D. Wu, "Machine Learning based Prediction Analysis of Potential Factors in Traffic Accidents," Applied and Computational Engineering, vol. 99, pp. 112–120, Nov. 2024.',
        '[3] Toronto Police Service, "Traffic collision data," City of Toronto Open Data Portal. https://open.toronto.ca/',
        '[4] Department for Transport (UK), "Road casualty statistics, collisions 2024," 2025.',
        '[5] scikit-learn developers, "scikit-learn 1.7.2 documentation," 2025. https://scikit-learn.org/stable/',
        '[6] PyTorch Team, "PyTorch 2.12 documentation," 2025. https://pytorch.org/docs/',
        '[7] Smart-Shield AI Capstone Repository (Team 2B), 2026. ' + GITHUB,
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Sheridan College — School of Applied Computing — INFO53883 AI/ML Capstone — Spring/Summer 2026"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Created: {path}")
