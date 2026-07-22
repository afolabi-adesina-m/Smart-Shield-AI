#!/usr/bin/env python3
"""
Regenerate all explanation artifacts from definitions.json + live project scan.

Outputs (explanations/):
  - Glossary.docx
  - Sprint-Pipeline-Swimlane.docx
  - Sprint-Pipeline-Swimlane.mmd
  - Train-Test-Validation-Splits.docx
  - README.md (index with last-built timestamp)
  - MANIFEST.json (source mtimes + output list)

Usage:
  python explanations/build_all.py          # one-shot build
  python explanations/build_all.py --watch  # rebuild when sources change
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

EXPLANATIONS = Path(__file__).resolve().parent
PROJECT = EXPLANATIONS.parent
DEFINITIONS_PATH = EXPLANATIONS / "definitions.json"

WATCH_PATHS = [
    DEFINITIONS_PATH,
    PROJECT / "notebooks" / "capstone_with_results.ipynb",
    PROJECT / "notebooks" / "capstone.ipynb",
    PROJECT / "src",
    PROJECT / "models" / "README.md",
    PROJECT / "Data" / "README.md",
]


def load_definitions() -> dict:
    return json.loads(DEFINITIONS_PATH.read_text(encoding="utf-8"))


def _mtime(path: Path) -> float | None:
    return path.stat().st_mtime if path.is_file() else None


def scan_notebook(nb_path: Path) -> dict:
    """Extract live metadata from the primary notebook."""
    meta: dict = {
        "path": str(nb_path.relative_to(PROJECT)) if nb_path.is_file() else None,
        "exists": nb_path.is_file(),
        "code_cells": 0,
        "markdown_cells": 0,
        "sections": [],
        "toronto_rows": None,
        "dft_rows": None,
    }
    if not nb_path.is_file():
        return meta

    nb = json.loads(nb_path.read_text(encoding="utf-8"))
    for cell in nb.get("cells", []):
        kind = cell.get("cell_type")
        if kind == "code":
            meta["code_cells"] += 1
            src = "".join(cell.get("source", []))
            for line in src.splitlines():
                if "df_toronto :" in line and "rows" in line:
                    m = re.search(r"([\d,]+)\s+rows", line)
                    if m:
                        meta["toronto_rows"] = m.group(1)
                if "dft        :" in line and "rows" in line:
                    m = re.search(r"([\d,]+)\s+rows", line)
                    if m:
                        meta["dft_rows"] = m.group(1)
        elif kind == "markdown":
            meta["markdown_cells"] += 1
            text = "".join(cell.get("source", [])).strip()
            first = text.split("\n")[0].strip()
            if re.match(r"#{1,3}\s+(Section|Sprint|##)", first) or "Section" in first[:40]:
                if len(first) < 120:
                    meta["sections"].append(first.lstrip("#").strip())

    return meta


def scan_data_rows() -> dict:
    """Quick row counts from CSV files if present."""
    stats = {"toronto_rows": None, "dft_rows": None}
    toronto = PROJECT / "Data" / "traffic collision data.csv"
    dft = PROJECT / "Data" / "dft-road-casualty-statistics-collision-2024.csv"
    try:
        if toronto.is_file():
            import pandas as pd
            stats["toronto_rows"] = f"{len(pd.read_csv(toronto, usecols=[0])):,}"
    except Exception:
        pass
    try:
        if dft.is_file():
            import pandas as pd
            stats["dft_rows"] = f"{len(pd.read_csv(dft, usecols=[0])):,}"
    except Exception:
        pass
    return stats


def compute_split_stats(defn: dict, nb_meta: dict, data_stats: dict) -> dict:
    splits = defn["splits"]
    test_size = splits["tabular_test_size"]
    train_frac = 1.0 - test_size

    toronto_s = (
        nb_meta.get("toronto_rows")
        or data_stats.get("toronto_rows")
        or "~809,034"
    )
    toronto_n = int(toronto_s.replace(",", ""))
    train_n = int(toronto_n * train_frac)
    test_n = toronto_n - train_n

    return {
        "test_size_pct": int(test_size * 100),
        "train_frac_pct": int(train_frac * 100),
        "toronto_rows": toronto_s,
        "train_rows": f"{train_n:,}",
        "test_rows": f"{test_n:,}",
        "cv_folds": splits["gridsearch_cv_folds"],
        "search_pct": int(splits["gridsearch_search_fraction"] * 100),
        "vision_train_pct": int(splits["vision_train_fraction"] * 100),
        "vision_val_pct": int(splits["vision_val_fraction"] * 100),
    }


def build_glossary(defn: dict, built_at: str) -> None:
    doc = Document()
    h = doc.add_heading("Smart-Shield Capstone Glossary", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = defn["project"]
    doc.add_paragraph(f"{p['course']} - {p['team']} - Sheridan College")
    doc.add_paragraph(
        f"Auto-generated {built_at} from explanations/definitions.json. "
        f"Primary notebook: {p['notebook_primary']}"
    )
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for term, definition in defn["glossary"]:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = definition
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)
    doc.save(EXPLANATIONS / "Glossary.docx")


def build_swimlane_doc(defn: dict, split: dict, built_at: str) -> None:
    doc = Document()
    h = doc.add_heading("Smart-Shield Pipeline Swimlane Diagram", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Auto-generated {built_at}")
    doc.add_heading("Sprint Overview", level=2)
    st = doc.add_table(rows=len(defn["sprint_overview"]) + 1, cols=3)
    st.style = "Table Grid"
    for j, hdr in enumerate(["Sprint", "Focus", "Notebook Sections"]):
        st.rows[0].cells[j].text = hdr
    for i, row in enumerate(defn["sprint_overview"], 1):
        for j, val in enumerate(row):
            st.rows[i].cells[j].text = val

    doc.add_heading("Swimlane: Step-by-Step Pipeline", level=2)
    for lane in defn["lanes"]:
        doc.add_heading(lane["name"], level=3)
        p = doc.add_paragraph()
        p.add_run(f"Sprint: {lane['sprint']}  |  Sections: {lane['sections']}").italic = True
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for j, hdr in enumerate(["Step", "Action", "What Happens", "Key Output"]):
            t.rows[0].cells[j].text = hdr
        for step in lane["steps"]:
            row = t.add_row()
            for j, val in enumerate(step):
                row.cells[j].text = val
        doc.add_paragraph()

    doc.add_heading("Train / Test Split (live)", level=2)
    sp = doc.add_table(rows=4, cols=3)
    sp.style = "Table Grid"
    sp.rows[0].cells[0].text = "Set"
    sp.rows[0].cells[1].text = "Share"
    sp.rows[0].cells[2].text = "Used For"
    rows = [
        ("Training", f"{split['train_frac_pct']}% ({split['train_rows']} rows)", "Model fitting + CV"),
        ("Test (held out)", f"{split['test_size_pct']}% ({split['test_rows']} rows)", "Final evaluation only"),
        ("Vision val", f"{split['vision_val_pct']}%", "ResNet18 early stopping / accuracy"),
    ]
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            sp.rows[i].cells[j].text = val

    doc.save(EXPLANATIONS / "Sprint-Pipeline-Swimlane.docx")


def build_swimlane_mmd(defn: dict) -> None:
    lines = [
        "%% Auto-generated by explanations/build_all.py — do not edit by hand",
        "%% Re-run: python explanations/build_all.py",
        "",
        "flowchart TB",
    ]
    lane_ids = []
    for i, lane in enumerate(defn["lanes"]):
        lid = f"L{i}"
        lane_ids.append(lid)
        lines.append(f'    subgraph {lid}["{lane["name"]}"]')
        lines.append("        direction LR")
        node_ids = []
        for step in lane["steps"]:
            sid = f"S{step[0]}"
            label = step[1].replace('"', "'")
            detail = step[2].replace('"', "'")[:60]
            lines.append(f'        {sid}["{sid}: {label}\\n{detail}"]')
            node_ids.append(sid)
        for a, b in zip(node_ids, node_ids[1:]):
            lines.append(f"        {a} --> {b}")
        lines.append("    end")
        lines.append("")

    for a, b in zip(lane_ids, lane_ids[1:]):
        lines.append(f"    {a} --> {b}")

    (EXPLANATIONS / "Sprint-Pipeline-Swimlane.mmd").write_text(
        "\n".join(lines) + "\n", encoding="utf-8"
    )


def build_splits_doc(defn: dict, split: dict, built_at: str) -> None:
    doc = Document()
    h = doc.add_heading("Training, Testing & Validation Splits", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Auto-generated {built_at}")
    doc.add_paragraph(
        f"Toronto dataset size used for estimates: {split['toronto_rows']} rows."
    )

    doc.add_heading("Tabular Models (Section 8)", level=2)
    doc.add_paragraph(
        f"train_test_split(test_size={defn['splits']['tabular_test_size']}, "
        f"random_state={defn['splits']['tabular_random_state']}, stratify=y)"
    )
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Set"
    t.rows[0].cells[1].text = "Share"
    t.rows[0].cells[2].text = "Purpose"
    for i, row in enumerate([
        ("Train", f"{split['train_frac_pct']}% ({split['train_rows']})", "Fit models; SMOTE applied here only"),
        ("Test", f"{split['test_size_pct']}% ({split['test_rows']})", "Held-out final evaluation"),
        ("CV (on train)", f"{split['cv_folds']}-fold StratifiedKFold", "Hyperparameter tuning inside GridSearchCV"),
    ], 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val

    doc.add_heading("GridSearchCV", level=2)
    doc.add_paragraph(
        f"Search runs on {split['search_pct']}% stratified sample of training data; "
        "best params refit on full training set."
    )

    doc.add_heading("Vision Model (Part 02)", level=2)
    vision_n = defn["splits"].get("vision_n_images", "curated RSCD sample")
    vision_bal = defn["splits"].get("vision_class_balance", "Clear / Wet / Snow balanced")
    doc.add_paragraph(
        f"Working sample: {vision_n} images ({vision_bal}). "
        f"{split['vision_train_pct']}% train / {split['vision_val_pct']}% validation. "
        "ResNet18 selected; autoencoder retained experimentally (use_hybrid=False)."
    )

    doc.add_heading("What Is NOT Split", level=2)
    for bullet in [
        "EDA (Sections 1-5): full datasets",
        "SHAP (Section 10.3): sample from test set",
        "Live demo scenarios: hand-crafted cases",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.save(EXPLANATIONS / "Train-Test-Validation-Splits.docx")


def build_readme(built_at: str, outputs: list[str]) -> None:
    text = f"""# Explanations

Auto-generated documentation for the Smart-Shield capstone.  
**Last built:** {built_at}  
**Regenerate:** `python explanations/build_all.py`

## Outputs

| File | Description |
|------|-------------|
| `Glossary.docx` | Terms and definitions (from `definitions.json`) |
| `Sprint-Pipeline-Swimlane.docx` | Step-by-step sprint pipeline tables |
| `Sprint-Pipeline-Swimlane.mmd` | Mermaid swimlane (render at [mermaid.live](https://mermaid.live)) |
| `Train-Test-Validation-Splits.docx` | Train/test/CV split reference |
| `MANIFEST.json` | Build metadata and source timestamps |
| `definitions.json` | **Edit this** to change glossary, lanes, split constants |

## Automatic updates

Explanations rebuild automatically when:

1. **Cursor hook** — saving `notebooks/*.ipynb` or `explanations/definitions.json` (`.cursor/hooks.json`)
2. **Git pre-commit** — run `git config core.hooksPath .githooks` once, then each commit rebuilds docs
3. **Manual** — double-click `explanations/update.bat` or run `python explanations/build_all.py`
4. **Watch mode** — `python explanations/build_all.py --watch` (polls every 30s)

## Customize content

Edit `definitions.json` then run `build_all.py`. Add glossary terms, lane steps, or change split constants (`tabular_test_size`, etc.).

## Files built this run

"""
    for name in outputs:
        text += f"- `{name}`\n"
    (EXPLANATIONS / "README.md").write_text(text, encoding="utf-8")


def write_manifest(built_at: str, outputs: list[str], sources: dict) -> None:
    manifest = {
        "generated_at": built_at,
        "generator": "explanations/build_all.py",
        "sources": sources,
        "outputs": outputs,
    }
    (EXPLANATIONS / "MANIFEST.json").write_text(
        json.dumps(manifest, indent=2), encoding="utf-8"
    )


def build_all(verbose: bool = True) -> list[str]:
    defn = load_definitions()
    nb_path = PROJECT / defn["project"]["notebook_primary"]
    nb_meta = scan_notebook(nb_path)
    data_stats = scan_data_rows()
    split = compute_split_stats(defn, nb_meta, data_stats)

    built_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    build_glossary(defn, built_at)
    build_swimlane_doc(defn, split, built_at)
    build_swimlane_mmd(defn)
    build_splits_doc(defn, split, built_at)

    outputs = [
        "Glossary.docx",
        "Sprint-Pipeline-Swimlane.docx",
        "Sprint-Pipeline-Swimlane.mmd",
        "Train-Test-Validation-Splits.docx",
        "README.md",
        "MANIFEST.json",
    ]

    sources = {
        str(p.relative_to(PROJECT)): _mtime(p)
        for p in WATCH_PATHS
        if p.exists()
    }
    sources["notebook_sections_found"] = len(nb_meta.get("sections", []))

    build_readme(built_at, outputs)
    write_manifest(built_at, outputs, sources)

    if verbose:
        print(f"[build_all] {built_at}")
        for name in outputs:
            print(f"  updated {name}")
        if nb_meta.get("toronto_rows"):
            print(f"  notebook toronto rows: {nb_meta['toronto_rows']}")
    return outputs


def watch_loop(interval: float = 30.0) -> None:
    print(f"Watching for changes (every {interval}s). Ctrl+C to stop.")
    last_sig: tuple | None = None
    while True:
        sig = tuple(_mtime(p) or 0 for p in WATCH_PATHS if p.exists())
        if sig != last_sig:
            if last_sig is not None:
                print("\n[watch] source changed — rebuilding...")
            build_all()
            last_sig = sig
        time.sleep(interval)


def main() -> int:
    parser = argparse.ArgumentParser(description="Rebuild Smart-Shield explanation docs")
    parser.add_argument("--watch", action="store_true", help="Poll sources and rebuild on change")
    parser.add_argument("--quiet", action="store_true")
    args = parser.parse_args()
    if args.watch:
        watch_loop()
        return 0
    build_all(verbose=not args.quiet)
    return 0


if __name__ == "__main__":
    sys.exit(main())
